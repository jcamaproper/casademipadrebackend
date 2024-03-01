from docx import Document
import requests
import re
from unidecode import unidecode
from app.insert_data import insertar_datos
#from firebase import send_push_notifications
import json
import os
import json


# Load environment variables from .env file

def analizar_documento(file_url, podcast_url):
    # Abre el documento de Word
    with requests.get(file_url, stream=True) as response:
        response.raise_for_status()  # This will raise an exception for HTTP errors

        # Create a temporary file
        temp_path = '/tmp/temp_document.docx'  # Adjust the path as needed

        # Stream the content to a file
        with open(temp_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192): 
                file.write(chunk)

    doc = Document(temp_path)
    map = {}
    parrafos = []
    # Se eliminan los espacios en blanco para no tener problemas con los saltos de línea
    for index, para in enumerate(doc.paragraphs):
        if para.text == "":
            continue
        else:
            parrafos.append(para.text)

    for index, para in enumerate(parrafos):
        #print(f"row {index}: {para}")
        if index == 0:
            map["semana"] = para
        elif index == 3:
            map["titulo_video"] = para
        elif index == 4:
            map["video_link"] = para
        elif index == 5:
            map["descripcion_video"] = para
        elif index == 6:
            map["titulo_audio"] = para
        elif index == 8:         
            map["descripcion_audio"] = para
        elif index == 9:
            map["soundcloud_link"] = para




    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                # Tabla 0, Titulo, fecha, tema, instrucciones
                if i==0 and j==0 and k==0: 
                    #print("Titulo:")
                    #print(cell.text)
                    map["titulo"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==0 and j==1 and k==0:
                    #print("Tema:")
                    #print(cell.text)
                    map["tema"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==0 and j==2 and k==0: 
                    #print("Instrucciones:")
                    #print(cell.text)
                    map["instrucciones"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                # Tabla 1, Devocional, Reflexion, Capitulo, *Biografia
                elif i==1 and j==3 and k==0: # Tabla 1, fila 3, celda 0
                    #print("Devocional:")
                    #print(cell.text)
                    map["devocional"] = cell.text
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==4 and k==0: # Tabla 1, fila 2, celda 0
                    #print("Reflexion:")
                    #print(cell.text)
                    map["reflexion"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==5 and k==0: 
                    #print("Capitulo:")
                    #print(cell.text)
                    map["capitulo"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==6 and k==0: 
                    #print("Lectura:")
                    #print(cell.text)
                    map["lectura"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

                elif i==1 and j==8 and k==0: 
                    #print("Biografia:")
                    #print(cell.text)
                    map["biografia"] = cell.text   
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")
                elif i==2 and j==1 and k==0:
                    #print("Trivia:")
                    #print(cell.text)
                    trivia= cell.text
                    map["trivia"] = extract_questions_to_map(trivia)
                    #print(f"Table {i}, Row {j}, Cell {k}: {cell.text}")

    #json_acentos = json.dumps(map, indent = 4)
    #json_sin_acentos = quitar_acentos_en_json(json_acentos)
    #print(json_sin_acentos)
    #print(map)

    map["podcast"] = podcast_url

    libro, personaje, biografia = extract_biography(map["biografia"])
    map["libro"] = libro
    map["personaje"] = personaje
    map["texto"] = biografia

    insertar_datos(map)

    #Despues de insertar los datos, se debe enviar una notificacion a los usuarios
    #send_push_notifications()
    os.remove(temp_path)
    return map


def quitar_acentos(texto):
    return unidecode(texto)


def quitar_acentos_en_json(json_data):
    data_dict = json.loads(json_data)
    for key, value in data_dict.items():
        if isinstance(value, str):
            data_dict[key] = quitar_acentos(value)
    return json.dumps(data_dict, indent=4)



def extract_questions_to_map(text):
    lines = text.strip().split('\n')
    questions_map = {}
    current_chapter = None

    # Process each line
    for line in lines:
        # Detect chapter lines by the absence of question marks and presence of numbers
        if '?' not in line and any(char.isdigit() for char in line):
            current_chapter = line.strip()
            questions_map[current_chapter] = []
        elif '?' in line:
            # Append the question to the list within the map under the current chapter
            if current_chapter:
                questions_map[current_chapter].append(line.strip())

    return questions_map

def extract_biography(text):
    # Extract 'libro'
    libro_match = re.search(r"Personajes principales de (\w+)", text)
    libro = libro_match.group(1) if libro_match else None

    # Extract 'personaje'
    personaje_match = re.search(r"\n(\w+)", text)
    personaje = personaje_match.group(1) if personaje_match else None

    # Extract 'biografia'
    biografia_start = text.find(personaje) + len(personaje)
    biografia = text[biografia_start:].strip()

    return libro, personaje, biografia

# Reemplaza 'ruta_del_documento.docx' con la ruta de tu propio documento de Word
# ruta_del_documento = 'template4.docx'
# if __name__ == "__main__":
#     analizar_documento(ruta_del_documento)

# def obtener_devocionales(offset=0, limite=10):
#     # Conexión a la base de datos
#     with get_db_cursor() as cur:
#         cur.execute("""
#             SELECT * FROM devocionales
#             ORDER BY fecha DESC
#             OFFSET %s LIMIT %s
#         """, (offset, limite))
#         registros = cur.fetchall()
#         columnas = [desc[0] for desc in cur.description]
#         devocionales = []
#         for registro in registros:
#             devocionales.append(dict(zip(columnas, registro)))
#     return devocionales

